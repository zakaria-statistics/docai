from pathlib import Path
from docx import Document as DocxDocument
from src.loaders.base_loader import BaseLoader
from src.models.document import Document


class DOCXLoader(BaseLoader):
    """Loader for Microsoft Word documents (.docx)."""

    def __init__(self, file_path: Path):
        super().__init__(file_path)

    def extract_text(self) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(self.file_path)
            paragraphs = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {e}")

    def load(self) -> Document:
        """Load the DOCX document."""
        content = self.extract_text()

        return Document.from_file(
            self.file_path,
            content,
            word_count=len(content.split()),
        )
